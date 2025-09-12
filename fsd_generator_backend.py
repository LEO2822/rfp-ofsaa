import numpy as np
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from qdrant_client import QdrantClient
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv
import os
import io
import logging
from datetime import datetime

# Load environment variables
load_dotenv()

# Configure logging for token usage
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TokenUsageTracker:
    """Track token usage and costs for FSD generation"""
    
    # GPT-4o-mini pricing (per 1M tokens)
    INPUT_COST_PER_1M = 0.15  # $0.15 per 1M input tokens
    OUTPUT_COST_PER_1M = 0.60  # $0.60 per 1M output tokens
    
    def __init__(self):
        self.total_input_tokens = 0
        self.total_output_tokens = 0
        self.total_cost = 0.0
        self.session_logs = []
    
    def log_usage(self, operation, input_tokens, output_tokens, context_info=""):
        """Log token usage for an operation"""
        input_cost = (input_tokens / 1_000_000) * self.INPUT_COST_PER_1M
        output_cost = (output_tokens / 1_000_000) * self.OUTPUT_COST_PER_1M
        total_cost = input_cost + output_cost
        
        self.total_input_tokens += input_tokens
        self.total_output_tokens += output_tokens
        self.total_cost += total_cost
        
        log_entry = {
            "timestamp": datetime.now().isoformat(),
            "operation": operation,
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "input_cost": round(input_cost, 6),
            "output_cost": round(output_cost, 6),
            "total_cost": round(total_cost, 6),
            "context_info": context_info
        }
        
        self.session_logs.append(log_entry)
        
        logger.info(f"🪙 TOKEN USAGE - {operation}")
        logger.info(f"   Input: {input_tokens:,} tokens (${input_cost:.6f})")
        logger.info(f"   Output: {output_tokens:,} tokens (${output_cost:.6f})")
        logger.info(f"   Total Cost: ${total_cost:.6f}")
        if context_info:
            logger.info(f"   Context: {context_info}")
        logger.info(f"   Session Total: ${self.total_cost:.6f}")
        
        return log_entry
    
    def get_session_summary(self):
        """Get summary of current session usage"""
        return {
            "total_input_tokens": self.total_input_tokens,
            "total_output_tokens": self.total_output_tokens,
            "total_tokens": self.total_input_tokens + self.total_output_tokens,
            "total_cost": round(self.total_cost, 6),
            "documents_generated": len([log for log in self.session_logs if log["operation"] == "FSD_Generation"]),
            "average_cost_per_document": round(self.total_cost / max(1, len([log for log in self.session_logs if log["operation"] == "FSD_Generation"])), 6)
        }

class FSDRequest(BaseModel):
    question: str

class FunctionDocumentGenerator:
    def __init__(self, qdrant_url=None, qdrant_api_key=None, collections=None):
        # Initialize token usage tracker
        self.token_tracker = TokenUsageTracker()
        if collections is None:
            self.collection_names = ["Flexcube_user_guide_14.x", "Flexcube_Userguide_12.x"]
        else:
            self.collection_names = collections
        
        # Initialize Qdrant client (optional - will work without it)
        try:
            if os.getenv("QDRANT_URL") and os.getenv("QDRANT_API_KEY"):
                self.qdrant_client = QdrantClient(
                    url=os.getenv("QDRANT_URL"),
                    api_key=os.getenv("QDRANT_API_KEY")
                )
            else:
                self.qdrant_client = None
        except Exception as e:
            print(f"Warning: Could not initialize Qdrant client: {e}")
            self.qdrant_client = None
        
        # Initialize OpenAI client
        openai_key = os.getenv("OPENAI_API_KEY")
        if not openai_key:
            raise EnvironmentError("Missing OPENAI_API_KEY in .env file")
        
        self.client = OpenAI(api_key=openai_key)
        
        # Load a model that generates 896-dimensional vectors (optional)
        self.model = None
        print("SentenceTransformer disabled for faster startup")

    def generate_embeddings(self, query):
        """Generate embeddings for the given query with padding to 896 dimensions"""
        if not self.model:
            return None
            
        embeddings = self.model.encode(query)
        
        # Pad or truncate to 896 dimensions
        if embeddings.shape[0] < 896:
            padded_embedding = np.pad(embeddings, (0, 896 - embeddings.shape[0]), mode='constant')
        else:
            padded_embedding = embeddings[:896]
        
        return padded_embedding.tolist()

    def search_vector_db(self, query, top_k=5):
        """Search vector database for relevant information"""
        if not self.qdrant_client or not self.model:
            return []
            
        query_embedding = self.generate_embeddings(query)
        if not query_embedding:
            return []
            
        results = []
        
        for collection in self.collection_names:
            try:
                search_result = self.qdrant_client.search(
                    collection_name=collection,
                    query_vector=query_embedding,
                    limit=top_k
                )
                results.extend(search_result)
            except Exception as e:
                print(f"Error searching collection {collection}: {e}")
        
        return results

    def get_mcp_context(self, function_requirement):
        """Get context from Context7 MCP server"""
        try:
            # MCP tool configuration (similar to mcp.py)
            mcp_tool = {
                "type": "mcp",
                "server_label": "context7",
                "server_url": "https://mcp.context7.com/mcp",
                "require_approval": "always"
            }
            
            # Create a focused prompt for MCP to get relevant documentation
            mcp_prompt = f"""
            Please search for official documentation and technical specifications related to the following functional requirement:
            
            Requirement: {function_requirement}
            
            Focus on:
            - Current system capabilities and features
            - Technical implementation details
            - Best practices and recommendations
            - Any relevant Oracle/banking domain expertise
            
            Provide detailed technical information that would help in creating a Functional Specification Document.
            """
            
            # Call OpenAI with MCP tool
            response = self.client.responses.create(
                model="gpt-4o-mini",
                tools=[mcp_tool],
                input=mcp_prompt
            )
            
            # Track initial MCP call tokens (estimated)
            mcp_input_tokens = len(mcp_prompt.split()) * 1.3  # Rough estimate
            mcp_output_tokens = 0
            
            # Handle MCP approval requests
            max_retries = 3
            retry_count = 0
            mcp_content = ""
            
            while retry_count < max_retries:
                # Check for MCP approval requests
                if any(ent.type == "mcp_approval_request" for ent in response.output):
                    approvals = [{
                        "type": "mcp_approval_response",
                        "approval_request_id": ent.id,
                        "approve": True
                    } for ent in response.output if ent.type == "mcp_approval_request"]
                    
                    response = self.client.responses.create(
                        model="gpt-4o-mini",
                        previous_response_id=response.id,
                        input=approvals,
                        tools=[mcp_tool],
                    )
                    continue
                
                # Extract MCP content
                for ent in response.output:
                    if ent.type == "mcp_call":
                        try:
                            mcp_content = getattr(ent, 'content', '') or str(ent.get('results', '')) or ''
                        except AttributeError:
                            mcp_content = ''
                    elif ent.type == "message":
                        for msg in ent.content:
                            mcp_content += msg.text.strip() + "\n"
                
                if mcp_content:
                    # Estimate output tokens from MCP content
                    mcp_output_tokens = len(mcp_content.split()) * 1.3
                    break
                    
                retry_count += 1
            
            # Log MCP token usage
            if mcp_content:
                self.token_tracker.log_usage(
                    "MCP_Context7_Query", 
                    int(mcp_input_tokens), 
                    int(mcp_output_tokens),
                    f"Retrieved {len(mcp_content)} chars of context"
                )
            
            return mcp_content
            
        except Exception as e:
            print(f"Warning: Could not retrieve MCP context: {e}")
            return ""

    def generate_document_with_llama(self, function_requirement, qdrant_context="", mcp_context=""):
        """Generate document using OpenAI GPT with both Qdrant and MCP context"""
        
        # Combine contexts
        combined_context = ""
        if qdrant_context:
            combined_context += f"Vector Search Context:\n{qdrant_context}\n\n"
        if mcp_context:
            combined_context += f"MCP Documentation Context:\n{mcp_context}\n\n"
        
        prompt = f"""
        Generate a comprehensive FSD (Functional Specification Document) based on the following input requirements:

        User Input Requirement: {function_requirement}
        
        {combined_context if combined_context else "Context: No additional context available"}

        Create a detailed document with the following structure:
        1. INTRODUCTION
           Brief overview of the document purpose and scope.
        
        2. REQUIREMENT OVERVIEW
           Clear statement of the business requirements and objectives.
        
        3. CURRENT FUNCTIONALITY
           Description of how the system currently handles these requirements.
        
        4. PROPOSED FUNCTIONAL APPROACH
           Detailed explanation of the proposed solution and implementation.

        The document should be professional, precise, and provide clear insights for implementation.
        Focus on these four main sections as they constitute the core content.
        
        Use the provided context information to enhance accuracy and technical depth.
        If context contains Oracle/banking domain information, incorporate relevant technical details.
        
        Note: Additional sections like Validations, Interface Impact, Migration Impact, etc. will be included in the final document template but do not need to be generated.
        """

        response = self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional FSD document specialist with expertise in Oracle banking solutions and technical documentation."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )

        # Log token usage from the completion
        if hasattr(response, 'usage') and response.usage:
            self.token_tracker.log_usage(
                "FSD_Generation",
                response.usage.prompt_tokens,
                response.usage.completion_tokens,
                f"Qdrant: {len(qdrant_context)} chars, MCP: {len(mcp_context)} chars"
            )
        else:
            # Estimate tokens if usage not available
            estimated_input = len(prompt.split()) * 1.3
            estimated_output = len(response.choices[0].message.content.split()) * 1.3
            self.token_tracker.log_usage(
                "FSD_Generation",
                int(estimated_input),
                int(estimated_output),
                f"Estimated - Qdrant: {len(qdrant_context)} chars, MCP: {len(mcp_context)} chars"
            )

        return response.choices[0].message.content

    def generate_function_document(self, function_requirement):
        """Main method to generate function-specific document using both Qdrant and MCP Context7"""
        # Step 1: Search vector database for relevant context (optional)
        qdrant_context = ""
        if self.qdrant_client:
            try:
                vector_search_results = self.search_vector_db(function_requirement)
                qdrant_context = "\n".join([
                    result.payload.get('text', '') 
                    for result in vector_search_results 
                    if result.payload and 'text' in result.payload
                ])
            except Exception as e:
                print(f"Warning: Could not search vector database: {e}")
        
        # Step 2: Get MCP Context7 documentation (optional)
        mcp_context = self.get_mcp_context(function_requirement)
        
        # Step 3: Generate document using OpenAI with both contexts
        generated_document = self.generate_document_with_llama(
            function_requirement, 
            qdrant_context,
            mcp_context
        )
        
        # Log session summary
        summary = self.token_tracker.get_session_summary()
        logger.info(f"📊 SESSION SUMMARY - Generated {summary['documents_generated']} documents")
        logger.info(f"   Total Tokens: {summary['total_tokens']:,} (Input: {summary['total_input_tokens']:,}, Output: {summary['total_output_tokens']:,})")
        logger.info(f"   Total Cost: ${summary['total_cost']:.6f}")
        logger.info(f"   Average per Document: ${summary['average_cost_per_document']:.6f}")
        
        return generated_document
    
    def add_bookmark(self, paragraph, bookmark_name):
        """Add a bookmark to a paragraph"""
        run = paragraph.add_run()
        tag = run._r
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), bookmark_name)
        tag.append(start)
        
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        end.set(qn('w:name'), bookmark_name)
        tag.append(end)

    def add_hyperlink(self, paragraph, text, bookmark_name):
        """Create a single-click hyperlink to a bookmark in the document"""
        from docx.oxml.shared import OxmlElement, qn
        
        # Create the w:hyperlink tag
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        hyperlink.set(qn('w:history'), '1')
        
        # Create a new run element
        new_run = OxmlElement('w:r')
        
        # Create run properties
        rPr = OxmlElement('w:rPr')
        
        # Add color (blue)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        # Add underline
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)
        
        # Add "hyperlink" style
        style = OxmlElement('w:rStyle')
        style.set(qn('w:val'), 'Hyperlink')
        rPr.append(style)
        
        new_run.append(rPr)
        
        # Add text
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        
        return hyperlink

    def add_dotted_toc_entry(self, doc, text, bookmark_name, page_number):
        """Add a Table of Contents entry with dotted line"""
        paragraph = doc.add_paragraph()
        
        # Create hyperlink to the bookmark
        self.add_hyperlink(paragraph, text, bookmark_name)
        
        # Calculate number of dots
        dots = "." * (60 - len(text))
        
        # Add dots
        paragraph.add_run(" " + dots + " ")
        
        # Add page number
        page_run = paragraph.add_run(str(page_number))
        page_run.bold = True

    def add_page_number(self, doc):
        """Add page numbers to the document footer"""
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add page number field
        run = paragraph.add_run()
        field_code = OxmlElement('w:fldChar')
        field_code.set(qn('w:fldCharType'), 'begin')

        instruction_text = OxmlElement('w:instrText')
        instruction_text.set(qn('xml:space'), 'preserve')
        instruction_text.text = 'PAGE'

        field_char = OxmlElement('w:fldChar')
        field_char.set(qn('w:fldCharType'), 'end')

        run._r.append(field_code)
        run._r.append(instruction_text)
        run._r.append(field_char)

    def save_as_word(self, text, logo_path=None, filename="fsd_document.docx"):
        """Create a Word document from the generated text and return it as bytes"""
        try:
            # Create a new Document
            doc = Document()
            
            # Add logo image if path is provided
            if logo_path and os.path.exists(logo_path):
                try:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(1.77), height=Inches(1.02))
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                    # Add bold title text below the image
                    title_paragraph = doc.add_paragraph()
                    title_run = title_paragraph.add_run("Functional Specification Document\n [Bank Name]")
                    title_run.bold = True
                    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  
                except Exception as e:
                    print(f"Warning: Could not add logo: {e}")
            
            # Add Table of Contents heading
            doc.add_heading('Table of Contents', level=1)
            
            # Define TOC items and their corresponding bookmarks
            toc_items = [
                ("1. Introduction", "intro_bookmark"),
                ("2. Requirement Overview", "req_overview_bookmark"),
                ("3. Current Functionality", "current_func_bookmark"),
                ("4. Proposed Functionalty", "proposed_func_bookmark"),
                ("5. Validations and Error Messages", "validations_bookmark"),
                ("6. Interface Impact", "interface_bookmark"),
                ("7. Migration Impact", "migration_bookmark"),
                ("8. Assumptions", "assumptions_bookmark"),
                ("9. RS-FS Traceability", "traceability_bookmark"),
                ("10. Open and Closed Queries", "queries_bookmark"),
                ("11. Annexure", "annexure_bookmark")
            ]
            
            # Add TOC entries with hyperlinks
            for i, (item, bookmark) in enumerate(toc_items):
                self.add_dotted_toc_entry(doc, item, bookmark, i+1)
            
            # Insert a page break after the Table of Contents
            doc.add_page_break()
            
            # Add page numbers to the document
            self.add_page_number(doc)
            
            # Parse the generated text and add core sections
            sections = {
                "1. INTRODUCTION": "",
                "2. REQUIREMENT OVERVIEW": "",
                "3. CURRENT FUNCTIONALITY": "",
                "4. PROPOSED FUNCTIONAL APPROACH": ""
            }
            
            # Process the generated text to extract sections
            current_section = None
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                
                # Check if line is a section header - improved matching
                for section_title in sections.keys():
                    section_name = section_title.split(". ")[1]  # e.g., "INTRODUCTION"
                    section_number = section_title.split(". ")[0]  # e.g., "1"
                    
                    # Match various formats: "1. INTRODUCTION", "INTRODUCTION", "1.INTRODUCTION", etc.
                    if (section_name.lower() in line.lower() and 
                        (section_number in line or line.strip().upper().startswith(section_name.upper()))):
                        current_section = section_title
                        break
            
                if current_section and line.lower() not in [s.lower() for s in sections.keys()]:
                    if not any(title.lower() in line.lower() for title in [t.split(". ")[1].lower() for t in sections.keys()]):
                        sections[current_section] += line + "\n"
            
            # Map section titles to bookmark names
            section_bookmarks = {
                "1. INTRODUCTION": "intro_bookmark",
                "2. REQUIREMENT OVERVIEW": "req_overview_bookmark",
                "3. CURRENT FUNCTIONALITY": "current_func_bookmark",
                "4. PROPOSED FUNCTIONAL APPROACH": "proposed_func_bookmark"
            }
            
            # Ensure all sections have some content - add fallback content if empty
            if not sections["1. INTRODUCTION"].strip():
                sections["1. INTRODUCTION"] = f"""This Functional Specification Document (FSD) outlines the requirements and proposed implementation approach for the requested functionality. The document serves as a comprehensive guide for development teams to understand the scope, current state, and proposed changes to the system.

This document addresses the following requirement: {function_requirement[:200]}..."""

            if not sections["2. REQUIREMENT OVERVIEW"].strip():
                sections["2. REQUIREMENT OVERVIEW"] = f"The business requirement centers around: {function_requirement}"

            # Add each section to the document with bookmarks
            for i, (section_title, content) in enumerate(sections.items(), 1):
                # Create a heading
                heading = doc.add_heading(f"{i}. {section_title.split('. ')[1].title()}", level=1)
                
                # Add bookmark to this heading
                self.add_bookmark(heading, section_bookmarks[section_title])
                
                # Add content
                if content.strip():
                    paragraph = doc.add_paragraph(content)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(6)
                    paragraph_format.space_after = Pt(6)
                    paragraph_format.line_spacing = 1.15
                else:
                    paragraph = doc.add_paragraph("Content to be added.")
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(6)
                    paragraph_format.space_after = Pt(6)
                    paragraph_format.line_spacing = 1.15
            
            # Add additional sections with bookmarks
            additional_sections = [
                ("5. Validations and Error Messages", "NA.", "validations_bookmark"),
                ("6. Interface Impact", "NA.", "interface_bookmark"),
                ("7. Migration Impact", "NA", "migration_bookmark"),
                ("8. Assumptions", "To be determined.", "assumptions_bookmark"),
                ("11. Annexure", "To be added as required.", "annexure_bookmark")
            ]
            
            for title, content, bookmark in additional_sections:
                heading = doc.add_heading(title, level=1)
                self.add_bookmark(heading, bookmark)
                
                paragraph = doc.add_paragraph(content)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(6)
                paragraph_format.space_after = Pt(6)
                paragraph_format.line_spacing = 1.15
            
            # Add RS-FS Traceability section with table and bookmark
            heading = doc.add_heading("9. RS-FS Traceability", level=1)
            self.add_bookmark(heading, "traceability_bookmark")
            
            # Create simple table for RS-FS Traceability
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            
            # Add headers manually
            try:
                headers = ["S. No.", "RS Section", "RS Section Description", "FS Section / Description"]
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
            except Exception as e:
                print(f"Warning: Issue with traceability table: {e}")
            
            # Add Open and Closed Queries section with table and bookmark
            heading = doc.add_heading("10. Open and Closed Queries", level=1)
            self.add_bookmark(heading, "queries_bookmark")
            
            # Create queries table with simple structure
            query_table = doc.add_table(rows=2, cols=6)
            query_table.style = 'Table Grid'
            
            # Add headers manually
            try:
                query_headers = ["Sr. No", "Issue Details", "Date Raised", "Clarification", "Raised By", "Current Status"]
                for i, header in enumerate(query_headers):
                    query_table.rows[0].cells[i].text = header
            except Exception as e:
                print(f"Warning: Issue with queries table: {e}")
            
            # Return the document object
            return doc
        except Exception as e:
            print(f"Error generating Word document: {e}")
            return None

# Global instance
doc_generator = FunctionDocumentGenerator()

def setup_fsd_generator_routes(app: FastAPI):
    """Add FSD Generator routes to the FastAPI app"""
    
    @app.post("/generate-fsd")
    async def generate_fsd_document(request: FSDRequest):
        try:
            if not request.question.strip():
                raise HTTPException(status_code=400, detail="Question cannot be empty")
            
            # Generate document content using the advanced logic
            generated_content = doc_generator.generate_function_document(request.question)
            
            # Create Word document with logo
            logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
            word_doc = doc_generator.save_as_word(generated_content, logo_path=logo_path)
            
            if not word_doc:
                raise HTTPException(status_code=500, detail="Failed to generate Word document")
            
            # Save document to BytesIO
            doc_buffer = io.BytesIO()
            word_doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Return as downloadable file
            return StreamingResponse(
                io.BytesIO(doc_buffer.read()),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=fsd_document.docx"}
            )

        except Exception as e:
            print(f"Error generating FSD document: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error generating FSD document: {str(e)}")